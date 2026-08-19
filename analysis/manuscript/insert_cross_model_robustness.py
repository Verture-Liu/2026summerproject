from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "latest version_PaleoRigor_v5_benchmark.docx"
OUTPUT = ROOT / "latest version_PaleoRigor_cross_model_robustness.docx"


RESULTS_TEXT = (
    "We next repeated the frozen v5 design with DeepSeek-V4-Pro to test whether the result depended on the original "
    "V4-Flash planner. The tasks, files, prompts, validation rules, scorer, local tools, call order, and 24-run sample per "
    "arm were unchanged. PaleoRigor passed 24 of 24 Pro runs (100%; Wilson 95% confidence interval, 86.2–100%), whereas "
    "the matched Pro model with the minimal workflow prompt passed 19 of 24 (79.2%; 59.5–90.8%). The paired difference "
    "was 20.8 percentage points. All five discordant pairs favored PaleoRigor and none favored the raw-model arm; the "
    "descriptive exact two-sided McNemar p value was 0.0625. The result met the preregistered engineering criterion of "
    "at least 22 PaleoRigor successes and more successes than the matched control. Together with the V4-Flash result, it "
    "shows that the tested control-layer benefit was retained across two model configurations from the same provider; it "
    "does not establish a ranking of models or providers."
)

METHODS_TEXT = (
    "For the cross-model robustness check, we reused the frozen v5 manifest and changed only the requested model from "
    "DeepSeek-V4-Flash to DeepSeek-V4-Pro. The same eight scenarios, three repeats, two arms, prompts, temperature (0), "
    "thinking mode, JSON response mode, validation and scoring rules, local execution environment, and fixed denominator "
    "were retained. The Pro-specific hypothesis and decision rule were committed before the first formal call. It was "
    "supported only if Pro/PaleoRigor achieved at least 22 of 24 strict successes and exceeded Pro/raw. All 48 Pro labels "
    "were independently recomputed from retained completions, decisions, workflows, validation records, and executions; "
    "all matched and no API call failed. The already observed Flash data were used only as context."
)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.style = paragraph.style
    inserted.add_run(text)
    return inserted


def find_paragraph(document: Document, startswith: str):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(startswith):
            return paragraph
    raise ValueError(f"Paragraph not found: {startswith}")


def main() -> None:
    document = Document(SOURCE)

    benchmark_end = find_paragraph(document, "V5 therefore met the release-specific reliability target")
    insert_after(benchmark_end, RESULTS_TEXT)

    methods_end = find_paragraph(document, "We report two-sided Wilson 95% confidence intervals")
    insert_after(methods_end, METHODS_TEXT)

    discussion = find_paragraph(document, "The first research question concerned release-level reliability")
    discussion.text = (
        discussion.text
        + " Repeating the frozen comparison with V4-Pro produced 24 of 24 PaleoRigor successes versus 19 of 24 for its "
        "matched raw-model control. The control layer therefore retained its release-level performance across two models "
        "from one provider, although this small robustness check cannot establish cross-provider generality."
    )

    limitations = find_paragraph(document, "The final benchmark was also small and depended on one model configuration")
    limitations.text = (
        "The final benchmark and cross-model check were small and used two model configurations from one provider in one "
        "local software environment. Each model comparison contained 24 runs per arm across eight scenario types, so the "
        "Wilson intervals remained wide. With V4-Flash, the 16.7-percentage-point paired advantage was not statistically "
        "significant (exact McNemar p = 0.219). With V4-Pro, the corresponding difference was 20.8 percentage points and "
        "all five discordant pairs favored PaleoRigor, but the exact test remained above 0.05 (p = 0.0625). These fixed-budget "
        "results support an engineering robustness criterion, not universal model superiority. Moreover, v3 and v4 both "
        "achieved only 75.0%; those rounds informed development and cannot be treated as independent support for v5. Future "
        "evaluations should preregister larger panels across independent providers, environments, genuinely unfamiliar errors, "
        "and scientific claims."
    )

    bundle = find_paragraph(document, "Each case was stored as a reproducibility package")
    bundle.text = bundle.text.replace(
        "the v3–v5 benchmark manifests, preregistrations, run-level scores, summary files, independent verification report, and the Figure 2 plotting script and source data",
        "the v3–v5 benchmark manifests, the preregistered V4-Pro robustness run, run-level scores, summary files, independent verification reports, and the Figure 2 plotting script and source data",
    )

    availability = find_paragraph(document, "All evaluation datasets are public records or small example files")
    availability.text = availability.text + (
        " The V4-Pro requests, raw completions, workflows or blocked decisions, execution records, scores, cross-model summary, "
        "and verification report are stored under analysis/benchmark_multimodel/v4_pro."
    )

    references_seen = False
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "References":
            references_seen = True
            continue
        if references_seen and paragraph.text.strip():
            paragraph.paragraph_format.space_after = Pt(2)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
