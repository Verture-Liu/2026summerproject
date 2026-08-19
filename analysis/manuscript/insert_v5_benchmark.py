#!/usr/bin/env python3
"""Insert the verified v5 benchmark figure and aligned text into the manuscript."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "latest version_PaleoRigor_logical_coherence.docx"
OUTPUT = ROOT / "latest version_PaleoRigor_v5_benchmark.docx"
FIGURE = (
    ROOT
    / "Final data"
    / "manuscript_figures"
    / "benchmark_v5"
    / "figure2_paleorigor_v5_benchmark.png"
)


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def replace_figure(paragraph) -> None:
    if not FIGURE.is_file():
        raise FileNotFoundError(FIGURE)
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(FIGURE), width=Inches(6.65))


def set_cell(table, row: int, column: int, text: str) -> None:
    table.cell(row, column).text = text


def main() -> None:
    document = Document(SOURCE)
    p = document.paragraphs

    edits = {
        7: (
            "Ancient microbial evidence can reveal the health, diet, and environments of past populations. However, this evidence is scarce, easily contaminated, and sensitive to mistakes in file selection and data handling. A standardized analysis assistant could reduce these errors and make each decision easier to inspect. Paleobiologists still lack a system that turns a research question into a reviewable analysis while keeping experts in control. We developed PaleoRigor, a human-in-the-loop system that can run only predefined, inspectable analysis modules, which we call skills. It links each requested analysis to named source files, records every data change, and asks an expert to review the proposed workflow before local execution. Two development-stage qualification rounds exposed systematic output-contract and over-planning errors. After these errors were corrected, the frozen final system completed 23 of 24 runs successfully in a new held-out evaluation (95.8%; Wilson 95% confidence interval, 79.8–99.3%), exceeding the prespecified threshold of 22 of 24. The same model with a minimal workflow prompt completed 19 of 24 runs (79.2%). PaleoRigor also matched read counts and compressed-file sizes for six public sequencing records, recorded the removal of 114 duplicate peptide sequences, and kept a clear source-audit boundary for a retraction-associated record. These results show that routine source checks and data changes can be made visible before fragile ancient evidence is interpreted. PaleoRigor does not authenticate ancient molecules, identify contamination sources, or reconstruct past microbial ecology. Instead, it gives paleobiologists an auditable starting point for deciding whether further biological analysis is justified."
        ),
        11: (
            "In the final held-out evaluation, PaleoRigor passed 23 of 24 runs (95.8%), compared with 19 of 24 (79.2%) for the raw-model control; all 12 boundary decisions were handled correctly."
        ),
        28: (
            "We converted this design into three quantitative questions. First, did the frozen final system exceed the prespecified strict-success threshold on unseen supported and boundary tasks, and how did it compare descriptively with the same model under a minimal prompt? Second, did local sequencing files agree with their intended public records at the tested metadata level? Third, did unsuitable requests stop before incompatible tools or unsupported claims reached local execution? Figure 2 answers the first and third questions, Figure 4 answers the second, and Figures 3 and 5 show how traceability and claim boundaries appear in concrete outputs."
        ),
        33: (
            "We evaluated six control targets (Table 2): workflow validity, file compatibility, source agreement, traceability, dependency behavior, and repeated-run or boundary control. Each target had an explicit pass rule. Inputs had to point to available files or earlier outputs; file formats had to match skills; and public-record agreement required Δreads = 0 and Rbytes = 1. Traceability required final files, intermediate files, and run records. Missing software had to be reported, supported runs had to pass workflow, execution, and output checks, and boundary requests had to stop with the prespecified reason before unsafe execution."
        ),
        40: (
            "PaleoRigor does not replace specialist bioinformatics pipelines; it adds controls before and around their execution. Table 3 therefore compares operational functions rather than speed or biological accuracy. The six functions are planning validity, recorded table curation, public-record agreement, claim boundaries, output traceability, and repeated-run reliability. The raw-model arm in Figure 2 is a matched planning control using the same model and task files with a minimal workflow prompt; it is not a comparison with a complete external bioinformatics system. The three case studies then evaluate source and transformation evidence."
        ),
        42: (
            "Table 3. Operational comparison of PaleoRigor with existing workflow practices. This is not a head-to-head speed or biological-accuracy benchmark. The measured-outcome column reports results obtained for PaleoRigor in this study. Figure 2 provides a matched planning-control comparison using the same model under a minimal prompt; the remaining comparison statements describe common differences in workflow control and were not measured as external software trials."
        ),
        43: (
            "This operational framing leads to a specific testable claim: the frozen PaleoRigor release should complete more than 90% of prespecified supported and boundary runs while keeping every proposed action inspectable. The versioned benchmark below tests that claim and shows the matched raw-model control. Figures 3–5 then test whether the resulting records explain a table change, establish public-record agreement, and preserve an interpretation boundary."
        ),
        44: "Final held-out evaluation exceeded the prespecified reliability threshold",
        45: (
            "We used two earlier held-out qualification rounds to identify release-blocking errors. Each round contained four supported workflows and four boundary requests, repeated three times in PaleoRigor and raw-model arms. PaleoRigor completed 18 of 24 runs in v3 and 18 of 24 in v4 (75.0% in each round; Figure 2A). The failures exposed narrow but systematic problems: declared output names did not always map to canonical skill outputs, common chart-name variants were rejected, and the planner sometimes added sample-sheet validation when the referenced FASTQ files were unavailable. We retained these failed rounds as development evidence and corrected the corresponding output contracts and planning rules before freezing v5."
        ),
        46: (
            "The v5 evaluation used new files, identifiers, values, and task wording. It again contained four supported and four boundary scenarios, with three repetitions per arm, giving 24 runs for PaleoRigor and 24 for the matched raw-model control. A run passed only if the response parsed, used all required and no forbidden skills, passed production validation, and completed local execution; a boundary run had to stop with the prespecified reason. PaleoRigor passed 23 of 24 runs (95.8%; Wilson 95% confidence interval, 79.8–99.3%), exceeding the prespecified threshold of at least 22 of 24. The raw-model control passed 19 of 24 (79.2%; 59.5–90.8%; Figure 2B)."
        ),
        47: (
            "The absolute paired difference was 16.7 percentage points in favor of PaleoRigor. Five paired runs were successful only for PaleoRigor and one only for the raw model; the exact two-sided McNemar test was not significant (p = 0.219). The comparison therefore describes this fixed benchmark rather than establishing a general model-level advantage. By task class, PaleoRigor passed 11 of 12 supported workflows (91.7%) and all 12 boundary decisions (100%), whereas the raw model passed 9 of 12 supported workflows (75.0%) and 10 of 12 boundary decisions (83.3%)."
        ),
        48: (
            "Scenario-level results show where these differences arose (Figure 2C). PaleoRigor achieved 3 of 3 successes for FASTA curation, paired FASTQ quality control, peptide-table processing, and each of the four boundary scenarios. It achieved 2 of 3 for metadata-to-sample-sheet preparation. In the single failed run, the model omitted the .tsv suffix from the registered upload reference; the production validator reported a missing uploaded input and prevented execution. The raw model was least reliable for FASTA curation and file-format mismatch decisions (1 of 3 each) and passed 2 of 3 peptide-table runs."
        ),
        49: (
            "V5 therefore met the release-specific reliability target while preserving a visible failure mode. The result does not erase the two 75.0% qualification rounds, prove universal reliability, or establish statistical superiority over the raw-model control. Instead, the version trajectory shows how failed tests exposed reproducible engineering defects, how those defects were corrected before a new freeze, and how the final release performed on a new fixed task set. Expert review remains necessary because the confidence intervals are wide and the benchmark cannot represent every archaeological data type or scientific claim."
        ),
        51: (
            "Figure 2. Versioned engineering benchmark and final held-out evaluation. A, Strict run-level success across v3, v4, and v5. V3 and v4 were retained as development-stage qualification rounds; each contained 24 PaleoRigor and 24 raw-model runs. The frozen v5 release used new files, values, identifiers, and wording. The dashed line marks the prespecified v5 threshold of more than 90%, operationalized as at least 22 of 24 successful PaleoRigor runs. B, Final v5 strict success for PaleoRigor and the same model under a minimal workflow prompt. Points show observed proportions and horizontal bars show two-sided Wilson 95% confidence intervals. The paired absolute difference was 16.7 percentage points; the exact two-sided McNemar p value was 0.219. C, V5 successes by scenario. Supported workflows comprised FASTA curation, paired FASTQ quality control, peptide-table processing, and sample-sheet preparation. Boundary decisions tested a file-format mismatch, a missing reference, an unsupported authenticity or contamination claim, and a missing paired-end mate. Each cell reports successful runs out of three."
        ),
        73: (
            "The first research question concerned release-level reliability. After two qualification rounds exposed repeatable implementation defects, the frozen v5 system passed 23 of 24 unseen runs and exceeded the prespecified threshold (Figure 2A,B). The lone failure was a malformed uploaded-file reference that the validator stopped before execution. This result supports reliability for the tested release and task set, not for every future request."
        ),
        75: (
            "The third question concerned scientific and input boundaries. In v5, PaleoRigor returned the correct blocked decision in all 12 boundary runs, including format mismatch, missing-reference, unsupported-claim, and missing-mate scenarios (Figure 2C). This is stronger than the earlier qualification behavior, but it does not remove the need for expert judgment. Figure 5 shows the intended division: the system reports reproducible file evidence, while an expert prevents that evidence from being promoted into biological validation. Fixed rules can constrain known errors; they cannot determine archaeological meaning or molecular authenticity."
        ),
        78: (
            "A larger set of negative controls is still needed. Retracted or disputed studies and low-biomass datasets affected by reagent contamination would provide useful tests (Eisenhofer et al., 2019; Salter et al., 2014). The current retraction-associated case shows that experts can set a clear interpretation boundary, while the v5 benchmark shows that predefined boundary categories can be enforced consistently. Future tests should ask whether that control generalizes to unfamiliar claims, mixed contamination sources, and incomplete archaeological metadata."
        ),
        82: (
            "PaleoRigor is still a prototype control layer, not a complete ancient-DNA or paleomicrobiome platform. The evidence supports a 95.8% strict-success rate for the frozen v5 task set, recorded table changes, and agreement between local files and public records. It does not support ancient-molecule authentication, identification of contamination sources, reconstruction of microbial ecology, or superiority to complete bioinformatics pipelines. Table 3 is an operational comparison; the raw-model arm isolates the effect of the control layer on planning and execution, not biological accuracy. The peptide case tests general table handling, and the retraction-associated case lies outside ancient-microbiome research. Larger ancient-data panels, controlled contamination mixtures, and matched workflow comparisons are needed before wider use."
        ),
        83: (
            "The final benchmark was also small and depended on one model configuration and one local software setup. V5 contained 24 runs per arm across eight scenario types, and its Wilson interval remained wide. The 16.7-percentage-point paired advantage over the raw-model control was not statistically significant (exact McNemar p = 0.219). Moreover, v3 and v4 both achieved only 75.0%; those rounds informed subsequent development and cannot be treated as independent support for the v5 estimate. Future evaluations should preregister larger panels, use several planning models and environments, and include genuinely unfamiliar errors and scientific claims."
        ),
        85: (
            "PaleoRigor provides paleobiologists with a reviewable path from an archaeological question and public data to a traceable local analytical starting point. The frozen final system passed 23 of 24 held-out runs (95.8%), including all 12 boundary decisions, while the matched raw-model control passed 19 of 24. Six public sequencing records matched repository read counts and compressed-file sizes, and 114 removed duplicate rows remained explicitly documented. These findings support PaleoRigor as an auditable evidence-preparation layer, not as an autonomous interpreter of ancient biology. Its broader value will depend on larger independent benchmarks, stronger contamination controls, ancient-DNA-specific checks, and community-maintained skills. With those additions, PaleoRigor could help paleobiologists examine fragile microbial evidence while keeping sources, transformations, uncertainty, and expert judgment visible."
        ),
        108: (
            "The versioned benchmark used a fixed paired design. Each round contained eight scenarios: four supported workflows and four boundary requests. Each scenario was repeated three times in both a PaleoRigor arm and a raw-model arm, producing 24 runs per arm and 48 API calls per round. The arms used the same model configuration and alternating call order. The raw-model arm received the workflow schema, uploaded-file summaries, and available-skill catalogue with a minimal planning prompt. The PaleoRigor arm additionally received the staged workflow, file-compatibility, prerequisite, output-contract, and scientific-boundary rules used by the application. Supported workflows executed locally; boundary requests were never executed."
        ),
        109: (
            "V3 and v4 were retained as failed qualification rounds after PaleoRigor achieved 18 of 24 strict successes in each. Their failures were used for development but not relabelled as confirmatory successes. Before v5, we corrected canonical output mapping, FASTA/FASTQ suffix handling, peptide-statistics output contracts, common peptide-chart aliases, and metadata-only sample-sheet planning. We then froze the v5 manifest, files, prompts, scorer, threshold, and code before any v5 API call. V5 used new files, identifiers, values, and wording. The sample size remained fixed, with no selective reruns or replacement of failures."
        ),
        110: (
            "For each run r and required criterion d, qᵣ,d was 1 when the run passed the stated rule and 0 when it failed. Supported-run criteria covered parsing, required and forbidden skills, production workflow validation, and local execution. Boundary runs required a valid blocked response with the prespecified reason code. A run passed only when every relevant rule passed:"
        ),
        114: (
            "We report two-sided Wilson 95% confidence intervals for the strict-success proportions. The prespecified v5 primary rule was at least 22 of 24 PaleoRigor successes, corresponding to more than 90%. The paired PaleoRigor–raw difference was summarized in percentage points and with an exact two-sided McNemar test; this secondary test was descriptive because 24 pairs provide limited power. Every stored v5 label was independently recomputed from the retained completion, workflow or blocked decision, production-validation result, and execution record. All 48 labels matched. Manifests, run-level scores, the summary JSON, and Figure 2 source data are stored under analysis/benchmark_v3, analysis/benchmark_v4, analysis/benchmark_v5, and Final data/manuscript_figures/benchmark_v5."
        ),
        127: (
            "The evaluation included three versioned engineering-benchmark rounds and three dataset-level case studies (Table 4). V3 and v4 were development-stage qualification rounds; v5 was the frozen final held-out evaluation. The peptide table tested recorded data changes. The six-record FASTQ panel tested agreement between local files and public records. The retraction-associated record tested whether file-level evidence stayed separate from a disputed biological claim. Measured results included strict workflow success, boundary handling, duplicate removal, table structure, read counts, compressed-file sizes, and descriptive quality-control values."
        ),
        132: (
            "Each case was stored as a reproducibility package with final outputs, intermediate files, and run records. Large public sequencing files were not copied into the repository. Their accessions and metadata allow readers to retrieve the same ENA/SRA records. The repository stores processed outputs, workflows, quality-control summaries, manifests, checksums, and plotting scripts. It also stores the v3–v5 benchmark manifests, preregistrations, run-level scores, summary files, independent verification report, and the Figure 2 plotting script and source data. Supplementary Tables S1–S3 link manuscript claims to these files."
        ),
        135: (
            "All evaluation datasets are public records or small example files in the project workspace. Table 4 lists their identity, input type, role, and measured results. Supplementary Table S1 links final outputs, intermediate files, and run records to repository folders. Supplementary Tables S2–S6 link claims, expert actions, errors, and software settings to their evidence. The v3 and v4 qualification records and the frozen v5 manifest, run-level table, summary JSON, verification report, plotting script, source-data CSV, and figure exports are stored under analysis/benchmark_v3, analysis/benchmark_v4, analysis/benchmark_v5, docs/science-superpowers, and Final data/manuscript_figures/benchmark_v5."
        ),
        156: (
            "All figures were produced from author-created vector schematics or Python-based plots of analytical outputs. The graphical abstract incorporates an openly licensed archaeological photograph with attribution. No figure was generated with a text-to-image model. Figure 2 was generated in Python from the retained v3–v5 benchmark summaries and run-level records. Its source-data CSV, plotting script, and editable SVG, PDF, PNG, and TIFF exports are preserved with the supplementary analysis materials."
        ),
    }

    for index, text in edits.items():
        set_text(p[index], text)

    replace_figure(p[50])

    # Table 3 (document table index 2): measured engineering-reliability row.
    set_cell(document.tables[2], 6, 1, "Final v5: PaleoRigor 23/24 strict successes (95.8%) versus 19/24 (79.2%) for the matched raw-model control; all 12 PaleoRigor boundary decisions passed.")
    set_cell(document.tables[2], 6, 2, "Shows release-level reliability and boundary control for the fixed task set, while the nonsignificant paired comparison and wide intervals limit generalization.")

    # Table 4 (document table index 3): engineering benchmark dataset row.
    set_cell(document.tables[3], 4, 1, "Versioned v3–v5 benchmark; each round used eight scenarios, three repeats, and matched PaleoRigor/raw-model arms (24 runs per arm).")
    set_cell(document.tables[3], 4, 2, "Test strict workflow success, local execution, input/scientific boundaries, and the effect of the PaleoRigor control layer under fixed model settings.")
    set_cell(document.tables[3], 4, 3, "V3: 18/24; v4: 18/24; frozen v5: PaleoRigor 23/24 (95.8%) and raw model 19/24 (79.2%). V5 boundary decisions: 12/12.")

    # Supplementary evidence tables.
    set_cell(document.tables[5], 4, 1, "Versioned v3–v5 manifests, preregistrations, raw completions, workflows or blocked decisions, validation/execution records, and independent scoring outputs.")
    set_cell(document.tables[5], 4, 2, "v5 run_level.csv, summary.json, verification_report.md, Figure 2 source data, and SVG/PDF/PNG/TIFF exports.")
    set_cell(document.tables[5], 4, 3, "V3–v5 records are under analysis/benchmark_v3, analysis/benchmark_v4, analysis/benchmark_v5, and Final data/manuscript_figures/benchmark_v5.")

    set_cell(document.tables[6], 6, 0, "The frozen final release exceeded the prespecified strict-success threshold, while earlier qualification failures remained visible.")
    set_cell(document.tables[6], 6, 1, "v3–v5 manifests and summaries; v5 run_level.csv, summary.json, verification_report.md, and retained execution records.")
    set_cell(document.tables[6], 6, 2, "Figure 2; Table 3")
    set_cell(document.tables[6], 6, 3, "analysis/benchmark_v3, analysis/benchmark_v4, analysis/benchmark_v5, and Figure 2 source-data folders")

    # Supplementary Table S1 benchmark evidence route.
    set_cell(document.tables[4], 10, 1, "analysis/benchmark_v3; analysis/benchmark_v4; analysis/benchmark_v5")
    set_cell(document.tables[4], 10, 2, "https://github.com/Verture-Liu/2026summerproject/tree/main/analysis/benchmark_v5")
    set_cell(document.tables[4], 10, 3, "Versioned manifests, preregistrations, run-level scores, summary JSON, verification report, and Figure 2 source data; complete raw run bundles are retained in the project archive.")

    document.core_properties.title = "PaleoRigor manuscript with versioned v5 benchmark"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
