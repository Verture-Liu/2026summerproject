from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from make_manuscript_docx import add_bullets, add_figure, add_para, set_styles


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ancient_research_agent_manuscript_polished.docx"
FIG_DIR = ROOT / "Final data" / "manuscript_figures"
TASK_DIR = ROOT / "Final data" / "task_grouped_figures"
RETRACT_DIR = ROOT / "Final data" / "retracted_case_figure"


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(
        "A skill-constrained natural-language research agent for reproducible ancient DNA and paleomicrobiome data reuse"
    )
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("1F2933")

    for text in [
        "Author names to be completed",
        "Affiliations to be completed",
        "# These authors contributed equally, if applicable.",
        "*Correspondence should be addressed to: [corresponding author email]",
        "Keywords: ancient DNA; paleomicrobiome; natural-language agent; reproducibility; FASTQ quality control; source-data audit",
    ]:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(3)


def add_refs(doc: Document) -> None:
    doc.add_heading("References", level=1)
    refs = [
        "1. Pääbo, S. Molecular cloning of Ancient Egyptian mummy DNA. Nature 314, 644–645 (1985).",
        "2. Dabney, J., Meyer, M. & Pääbo, S. Ancient DNA damage. Cold Spring Harb. Perspect. Biol. 5, a012567 (2013).",
        "3. Warinner, C., et al. A robust framework for microbial archaeology. Annu. Rev. Genomics Hum. Genet. 18, 321–356 (2017).",
        "4. Andrews, S. FastQC: a quality control tool for high throughput sequence data (2010).",
        "5. Ewels, P., Magnusson, M., Lundin, S. & Käller, M. MultiQC: summarize analysis results for multiple tools and samples in a single report. Bioinformatics 32, 3047–3048 (2016).",
        "6. Shen, W., Le, S., Li, Y. & Hu, F. SeqKit: a cross-platform and ultrafast toolkit for FASTA/Q file manipulation. PLoS ONE 11, e0163962 (2016).",
        "7. Leinonen, R., Sugawara, H. & Shumway, M. The Sequence Read Archive. Nucleic Acids Res. 39, D19–D21 (2011).",
        "8. Harrison, P. W., et al. The European Nucleotide Archive in 2024. Nucleic Acids Res. 52, D92–D97 (2024).",
        "9. Schubert, M., Lindgreen, S. & Orlando, L. AdapterRemoval v2: rapid adapter trimming, identification, and read merging. BMC Res. Notes 9, 88 (2016).",
        "10. Jónsson, H., Ginolhac, A., Schubert, M., Johnson, P. L. F. & Orlando, L. mapDamage2.0: fast approximate Bayesian estimates of ancient DNA damage parameters. Bioinformatics 29, 1682–1684 (2013).",
        "11. Gao, W., et al. Predictable regulation of gut microbiome in immunotherapeutic efficacy of gastric cancer. Genes Immun. 26, 1–8 (2025). [Retracted]",
        "12. Gao, W., et al. Retraction Note: Predictable regulation of gut microbiome in immunotherapeutic efficacy of gastric cancer. Genes Immun. 27, 384 (2026).",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_styles(doc)

    add_title_block(doc)

    doc.add_heading("Abstract", level=2)
    add_para(
        doc,
        "Reusing public ancient DNA and paleomicrobiome data remains difficult because early processing steps require file-format inspection, command-line tools, repository metadata checks, and careful audit trails. We developed a skill-constrained natural-language research agent that converts user requests into validated local workflows. The language model plans the workflow, whereas predefined skills and external tools perform execution, including CSV curation modules, FastQC, MultiQC, and seqkit. We evaluated the prototype across three source-data reuse scenarios: peptide validation table curation, FASTQ metadata reproduction against ENA/SRA records, and audit of a retraction-associated sequencing record. The agent generated organized outputs, reproduced public read-count and file-size metadata, and preserved intermediate records for later inspection. These results suggest that constrained natural-language agents can lower the operational barrier to ancient-data reuse while maintaining transparent and auditable execution.",
    )

    doc.add_heading("Key points", level=2)
    add_bullets(
        doc,
        [
            "The agent converts natural-language requests into structured, validated, and locally executable workflows.",
            "Execution is constrained by predefined skills and standard tools, reducing the risk of arbitrary code generation.",
            "Validation covered peptide CSV curation, FASTQ source-data quality control, and audit of a retraction-associated sequencing record.",
            "The prototype targets entry-level but failure-prone steps in ancient DNA and paleomicrobiome data reuse.",
        ],
    )

    doc.add_heading("Graphical abstract", level=2)
    add_figure(
        doc,
        FIG_DIR / "graphical_abstract_three_scenarios.png",
        "Graphical abstract. The agent supports three practical scenarios in ancient human and paleomicrobiome data reuse: routine reproducibility, source-data checking, and audit of problematic or retraction-associated records.",
        width=6.5,
    )

    doc.add_heading("Introduction", level=2)
    add_para(
        doc,
        "Ancient human, ancient DNA, and paleomicrobiome studies increasingly depend on public data reuse. Sequencing reads, accession records, and supplementary tables often make reuse possible in principle. In practice, however, researchers must still identify the correct source files, inspect file formats, run command-line tools, compare local files with repository metadata, and retain intermediate outputs. The bottleneck is therefore not only data availability, but the translation of heterogeneous public records into reproducible local analyses.",
    )
    add_para(
        doc,
        "This bottleneck is especially consequential for ancient-data studies. Ancient DNA and paleomicrobiome datasets are often fragmented, low-input, degraded, and contamination-sensitive. Early checks such as FASTQ integrity, paired-end consistency, read-count agreement, read-length summaries, GC content, and malformed-record detection determine whether downstream alignment, damage profiling, taxonomic classification, and contamination assessment can be interpreted reliably.",
    )
    add_para(
        doc,
        "Large language models provide a natural interface for lowering the technical barrier to data processing. Yet unrestricted code generation is poorly suited to reproducible research workflows. Generated steps may reference missing files, assume unavailable tools, or construct incompatible workflows. These risks are amplified when users are new to command-line bioinformatics or when analyses depend on strict source-data provenance.",
    )
    add_para(
        doc,
        "We therefore designed a skill-constrained research agent for ancient-data reuse. The language model translates the user request into a structured workflow, but execution is delegated to predefined local skills and external tools. A validator checks file types, step references, and skill compatibility before execution. The system then exports final outputs, intermediate files, manifests, checksums, and reports. This design creates a controlled bridge between natural-language interaction and auditable local analysis.",
    )

    add_figure(
        doc,
        FIG_DIR / "figure1_agent_architecture_workflow.png",
        "Figure 1. System architecture of the natural-language research agent. User requests are translated into structured workflows, checked by a validator, executed by predefined local skills and command-line tools, and exported as final outputs, step outputs, and audit records.",
        width=6.5,
    )

    doc.add_heading("Results", level=2)
    doc.add_heading("Task 1: The agent curated a headerless peptide validation table", level=3)
    add_para(
        doc,
        "We first tested whether the agent could curate a small tabular biological dataset from a natural-language instruction. The input was a headerless peptide validation CSV containing 5,810 rows. The planned workflow normalized the table into canonical label and sequence columns, validated peptide sequences, removed duplicate entries, computed descriptive summaries, and exported a cleaned CSV file.",
    )
    add_para(
        doc,
        "The workflow produced a cleaned table with 5,696 rows and 5,696 unique sequences. In total, 114 duplicate peptide records were removed. The output preserved the expected label structure and peptide-length distribution, indicating that the agent performed the intended curation rather than altering the biological content arbitrarily. This task demonstrates that the agent can convert ambiguous tabular inputs into standardized, reusable validation data.",
    )
    add_figure(
        doc,
        TASK_DIR / "task1_peptide_csv_curation.png",
        "Figure 2. Peptide CSV curation. A, Workflow summary for headerless peptide-table processing. B, Row-count change after duplicate removal. C, Label balance before and after cleaning. D, Peptide-length distribution after curation.",
        width=6.5,
    )

    doc.add_heading("Task 2: The agent reproduced public FASTQ source metadata", level=3)
    add_para(
        doc,
        "We next evaluated FASTQ source-data quality control using six public sequencing records. The agent ran file inspection, FastQC, MultiQC, and seqkit-based statistics, then compared local outputs with ENA/SRA metadata. We treated read count and compressed file size as source-level agreement metrics. GC content and mean read length were reported as derived quality-control summaries, not as criteria for reproducing a biological conclusion.",
    )
    add_para(
        doc,
        "Across all six FASTQ records, agent-derived read counts matched the corresponding ENA/SRA records, with a read-count deviation of zero in each case. Local file sizes also matched public fastq_bytes values, yielding a local-to-public byte ratio of 1.000 for all tested records. The tested runs ranged from 56,122 to 4,380,359 reads, indicating that the same audit logic applied across small and larger source files.",
    )
    add_para(
        doc,
        "These results show that the agent can verify whether local sequencing inputs correspond to the intended public records before downstream analysis. For ancient DNA and paleomicrobiome reuse, this type of source-data confirmation is a practical safeguard against file mix-ups, incomplete downloads, and paired-end mismatches.",
    )
    add_figure(
        doc,
        TASK_DIR / "task2_fastq_source_qc.png",
        "Figure 3. FASTQ source-data quality control. A, Agent-derived read counts compared with public ENA read counts. B, Local compressed file sizes compared with public fastq_bytes values. C, GC content reported as a derived quality-control metric. D, Public read counts and agent-derived values overlap for the tested records.",
        width=6.5,
    )

    doc.add_heading("Task 3: The agent audited a retraction-associated public sequencing record", level=3)
    add_para(
        doc,
        "Finally, we tested whether the same framework could audit a source record associated with problematic literature. We selected the paired FASTQ run SRR29088443 from the public study SRP508771, which is linked to a retracted publication. This dataset was used as a negative-control audit case. The goal was not to evaluate the retracted biological claim, but to test whether the agent could reproduce and document source-level metadata from a problematic literature record.",
    )
    add_para(
        doc,
        "The agent identified the paired FASTQ files and reproduced the public metadata. Each mate contained 53,571 reads, matching the ENA/SRA record. The total base count was 26,517,645 bp, and the compressed file sizes matched the public record for both R1 and R2. The agent also reported derived quality-control metrics, including mean read lengths of 244 bp and 251 bp, GC contents of 54.0% and 53.6%, and zero malformed records.",
    )
    add_para(
        doc,
        "This task illustrates the boundary of the system. The agent can audit whether public source data are accessible, internally consistent, and processed transparently. It does not determine why a paper was retracted and does not validate the original biological conclusion.",
    )
    add_figure(
        doc,
        RETRACT_DIR / "task3_retracted_fastq_audit.png",
        "Figure 4. Task 3: source-data audit of a retraction-associated FASTQ record. A, Case information linking the public FASTQ run to the retracted publication and repository record. B, Deviation between agent-derived and ENA/SRA source-level metadata. C, Derived FASTQ quality metrics. D, Audit checklist distinguishing source-data verification from biological claim validation.",
        width=6.5,
    )

    doc.add_heading("Discussion", level=2)
    add_para(
        doc,
        "This study presents a prototype natural-language agent for reproducible ancient-data reuse. Its contribution is not a new ancient-DNA algorithm, but a constrained execution layer that helps users move from natural-language intent to inspectable local outputs. The validation tasks show that the agent can curate tabular biological data, reproduce public FASTQ source metadata, and audit a retraction-associated sequencing record while preserving intermediate evidence.",
    )
    add_para(
        doc,
        "The design differs from a general-purpose chatbot. The language model plans the workflow, but it does not freely execute arbitrary generated code. Instead, the workflow must pass validation and is executed through predefined skills and configured local tools. This separation reduces the risk of hallucinated operations and improves traceability. It also makes failures more interpretable, because missing dependencies or incompatible inputs can be reported before or during execution.",
    )
    add_para(
        doc,
        "For ancient DNA and paleomicrobiome studies, this front-end layer addresses a practical source of error. Before authentication, mapping, taxonomic profiling, or contamination analysis, researchers must confirm that source files are correct and traceable. The current system should therefore be viewed as an auditable front end for data reuse rather than an autonomous interpreter of ancient-DNA evidence.",
    )
    add_para(
        doc,
        "Several limitations remain. The current prototype focuses on entry-level source-data processing rather than complete ancient-DNA pipelines. Downstream analyses such as read alignment, damage-pattern estimation, contamination assessment, metagenomic classification, and authentication require additional skills and benchmark datasets. Workflow quality also depends on planner instructions, validator coverage, and the availability of external tools in the local environment.",
    )
    add_para(
        doc,
        "Future development will extend the skill library to ancient-DNA-specific analysis modules, including BWA or Bowtie2 alignment, samtools statistics, mapDamage profiling, contamination screening, and taxonomic profiling. A larger benchmark across published paleogenomic and paleomicrobiome datasets will be needed to evaluate robustness across data types, repository conventions, and laboratory contexts.",
    )

    doc.add_heading("Methods", level=2)
    doc.add_heading("Workflow schema and planning", level=3)
    add_para(
        doc,
        "The agent accepts uploaded files and a natural-language task description. The planner converts the request into a structured workflow containing ordered steps. Each step specifies a skill, input references, parameters, expected outputs, and a short reason. This explicit representation allows the workflow to be inspected and validated before execution.",
    )
    doc.add_heading("Validation and execution control", level=3)
    add_para(
        doc,
        "The validator checks step references, expected outputs, file-format compatibility, and skill availability. This step is intended to catch common planning errors, such as referencing missing outputs or applying a FASTQ-specific skill to a CSV input. Validated steps are dispatched to local skills, which call deterministic Python modules or configured external tools.",
    )
    doc.add_heading("Implemented skills and external tools", level=3)
    add_para(
        doc,
        "Implemented skills include peptide CSV normalization, peptide sequence validation, duplicate removal, file-type detection, FASTQ paired-end matching, FastQC execution, MultiQC report generation, seqkit statistics, and source-data audit utilities. External tools are run locally. When a required tool is not available, the agent reports the missing dependency rather than installing software automatically.",
    )
    doc.add_heading("Output organization", level=3)
    add_para(
        doc,
        "Each run generates an independent result directory. User-facing files are exported to final_outputs. Intermediate step files are stored in step_outputs. Workflow JSON, manifests, checksums, timestamps, and a run report are preserved in ResearchAgent Records. This separation allows users to retrieve final deliverables while retaining the evidence needed for reproducibility checks.",
    )
    doc.add_heading("Validation metrics", level=3)
    add_para(
        doc,
        "For peptide CSV curation, the primary metrics were input row count, cleaned row count, duplicate count, unique sequence count, label distribution, and peptide-length distribution. For FASTQ source-data QC, source-level agreement was evaluated using read-count differences and compressed file-size ratios relative to ENA/SRA metadata. GC content, mean read length, N fraction, and malformed-record counts were reported as quality-control summaries.",
    )
    doc.add_heading("Retraction-associated audit case", level=3)
    add_para(
        doc,
        "The retraction-associated case used SRR29088443 from SRP508771. The audit compared agent-derived records, base counts, and compressed file sizes with public repository metadata. The analysis was intentionally limited to source-data verification and QC report generation. It was not designed to evaluate the biological interpretation of the retracted article.",
    )

    doc.add_heading("Data Availability", level=2)
    add_para(
        doc,
        "The validation datasets are public or organized within the project workspace. Public sequencing records are linked to ENA/SRA accessions described in the Results and Methods. A final accession table and download commands should be added before submission.",
    )
    doc.add_heading("Code Availability", level=2)
    add_para(
        doc,
        "The Research Agent code is currently maintained in the local project repository. A public GitHub repository, release tag, installation guide, example data, and archived version should be provided before publication.",
    )
    doc.add_heading("Acknowledgments", level=2)
    add_para(doc, "Acknowledgments to supervisors, collaborators, and funding sources should be completed by the author.")
    doc.add_heading("Author contributions", level=2)
    add_para(doc, "Author contribution statements should be completed after the final author list is confirmed.")
    doc.add_heading("Competing interest", level=2)
    add_para(doc, "The authors declare no competing interests.")
    add_refs(doc)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
