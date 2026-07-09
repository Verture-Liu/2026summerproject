from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from make_manuscript_docx import add_bullets, add_figure, add_para, set_cell_shading, set_cell_text, set_styles


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ancient_research_agent_manuscript_v2_human_in_loop.docx"
FIG_DIR = ROOT / "Final data" / "manuscript_figures"
TASK_DIR = ROOT / "Final data" / "task_grouped_figures"
RETRACT_DIR = ROOT / "Final data" / "retracted_case_figure"


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(
        "A skill-constrained human-in-the-loop research agent for rigorous paleomicrobiome data analysis"
    )
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("1F2933")

    for text in [
        "Author names to be completed",
        "Affiliations to be completed",
        "# These authors contributed equally, if applicable.",
        "*Correspondence should be addressed to: [corresponding author email]",
        "Keywords: paleomicrobiome; ancient DNA; human-in-the-loop agent; workflow validation; FASTQ quality control; reproducible bioinformatics",
    ]:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(3)


def add_formula(doc: Document, formula: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(formula)
    run.font.name = "Cambria Math"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string("1F2933")


def add_compact_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    font_size: float = 9,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, text, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1.5)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph()


def add_refs(doc: Document) -> None:
    doc.add_heading("References", level=1)
    refs = [
        "1. Pääbo, S. Molecular cloning of Ancient Egyptian mummy DNA. Nature 314, 644–645 (1985). https://doi.org/10.1038/314644a0",
        "2. Dabney, J., Meyer, M. & Pääbo, S. Ancient DNA damage. Cold Spring Harb. Perspect. Biol. 5, a012567 (2013). https://doi.org/10.1101/cshperspect.a012567",
        "3. Warinner, C., et al. A robust framework for microbial archaeology. Annu. Rev. Genomics Hum. Genet. 18, 321–356 (2017). https://doi.org/10.1146/annurev-genom-091416-035526",
        "4. Andrews, S. FastQC: a quality control tool for high throughput sequence data. Babraham Bioinformatics. https://www.bioinformatics.babraham.ac.uk/projects/fastqc/ (accessed 6 July 2026).",
        "5. Ewels, P., Magnusson, M., Lundin, S. & Käller, M. MultiQC: summarize analysis results for multiple tools and samples in a single report. Bioinformatics 32, 3047–3048 (2016). https://doi.org/10.1093/bioinformatics/btw354",
        "6. Shen, W., Le, S., Li, Y. & Hu, F. SeqKit: a cross-platform and ultrafast toolkit for FASTA/Q file manipulation. PLoS ONE 11, e0163962 (2016). https://doi.org/10.1371/journal.pone.0163962",
        "7. Leinonen, R., Sugawara, H. & Shumway, M. The Sequence Read Archive. Nucleic Acids Res. 39, D19–D21 (2011). https://doi.org/10.1093/nar/gkq1019",
        "8. O’Cathail, C., et al. The European Nucleotide Archive in 2024. Nucleic Acids Res. 53, D49–D55 (2025). https://doi.org/10.1093/nar/gkae975",
        "9. Schubert, M., Lindgreen, S. & Orlando, L. AdapterRemoval v2: rapid adapter trimming, identification, and read merging. BMC Res. Notes 9, 88 (2016). https://doi.org/10.1186/s13104-016-1900-2",
        "10. Jónsson, H., Ginolhac, A., Schubert, M., Johnson, P. L. F. & Orlando, L. mapDamage2.0: fast approximate Bayesian estimates of ancient DNA damage parameters. Bioinformatics 29, 1682–1684 (2013). https://doi.org/10.1093/bioinformatics/btt193",
        "11. Gao, W., Wang, Z., Shi, X., Wu, H., Zhou, B., Lin, X., et al. RETRACTED ARTICLE: Predictable regulation of gut microbiome in immunotherapeutic efficacy of gastric cancer. Genes Immun. 26, 1–8 (2024). https://doi.org/10.1038/s41435-024-00306-2",
        "12. Gao, W., Wang, Z., Shi, X., Wu, H., Zhou, B., Lin, X., et al. Retraction Note: Predictable regulation of gut microbiome in immunotherapeutic efficacy of gastric cancer. Genes Immun. 27, 384 (2026). https://doi.org/10.1038/s41435-026-00397-z",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_styles(doc)

    add_title_block(doc)

    doc.add_heading("Abstract", level=2)
    add_para(
        doc,
        "Ancient human microbiome and paleomicrobiome studies require early-stage analyses that are technically simple in appearance but scientifically consequential in practice, including file-format inspection, source-record verification, FASTQ quality control, paired-end checking, and traceable output management. These steps are difficult for non-specialist users because they require command-line tools, repository metadata interpretation, and careful provenance control. We developed a skill-constrained human-in-the-loop research agent that converts natural-language requests into validated local workflows while preserving expert review points before execution and during result interpretation. The language model proposes the workflow, whereas predefined skills, validators, and local bioinformatics tools perform constrained execution. We evaluated the prototype at both system and dataset levels: workflow validation was used to prevent incompatible step references, and case studies tested peptide table curation, public FASTQ source-data quality control, and audit of a retraction-associated sequencing record. The agent reproduced source-level read counts and file-size metadata across six public FASTQ records, curated a 5,810-row peptide validation table into 5,696 non-redundant records, and produced an auditable source-data report for SRR29088443 without claiming to validate the retracted biological conclusion. These results indicate that natural-language agents can support rigorous paleomicrobiome data analysis when planning, execution, and expert interpretation are separated by explicit workflow constraints.",
    )

    doc.add_heading("Key points", level=2)
    add_bullets(
        doc,
        [
            "The system is framed as an auditable assistant for rigorous ancient human microbiome and paleomicrobiome analysis, not merely as a public-data reuse shortcut.",
            "The language model proposes workflows, but predefined skills, validators, local tools, and expert checkpoints control execution.",
            "Engineering evaluation focused on workflow validity, file-format compatibility, source-metadata agreement, output traceability, and missing-dependency behavior.",
            "Dataset-level demonstrations included peptide CSV curation, FASTQ source-data QC, and a retraction-associated source-data audit.",
        ],
    )

    doc.add_heading("Graphical abstract", level=2)
    add_figure(
        doc,
        FIG_DIR / "graphical_abstract_three_scenarios.png",
        "Graphical abstract. The agent is designed as a practical interface between expert intent and auditable local analysis for ancient human microbiome and paleomicrobiome data.",
        width=6.5,
    )

    doc.add_heading("Introduction", level=2)
    add_para(
        doc,
        "Ancient human microbiome and paleomicrobiome research increasingly depends on heterogeneous public data, including sequencing reads, repository accession records, and supplementary biological tables. Although these resources are often available in principle, rigorous analysis remains difficult in practice. Researchers must identify the correct files, confirm that local downloads match public records, inspect file formats, run specialized tools, and retain intermediate outputs for later review. The practical challenge is therefore not only access to data, but the ability to transform heterogeneous source material into controlled and auditable analyses.",
    )
    add_para(
        doc,
        "This requirement is especially important for ancient-data research. Ancient DNA and paleomicrobiome materials are commonly low-input, degraded, fragmented, contamination-sensitive, and dependent on strict provenance. Early failures, such as mismatched paired-end files, incomplete downloads, malformed FASTQ records, or unreported read-count differences, can compromise downstream mapping, damage profiling, taxonomic assignment, contamination assessment, and authentication. These checks require professional judgment even when the commands themselves are routine.",
    )
    add_para(
        doc,
        "Large language models offer a promising natural-language interface for lowering operational barriers, but unrestricted code generation is not sufficient for reproducible scientific analysis. A model may invent file names, reference outputs that do not exist, apply FASTQ tools to CSV files, or assume unavailable environments. In ancient-data workflows, these errors are not merely inconvenient; they can obscure whether a downstream result reflects biology, file handling, or a planning mistake.",
    )
    add_para(
        doc,
        "We therefore developed a skill-constrained human-in-the-loop research agent for rigorous paleomicrobiome data analysis. The system uses the language model for workflow proposal, but delegates execution to predefined skills and local tools after validation. Human experts remain involved at three points: reviewing the generated workflow, approving local execution, and interpreting final reports. This design is intended to preserve the flexibility of natural-language interaction while maintaining explicit controls around data provenance, file compatibility, and output traceability.",
    )

    add_figure(
        doc,
        FIG_DIR / "figure1_agent_architecture_workflow.png",
        "Figure 1. Human-in-the-loop architecture of the research agent. User requests are converted into structured workflows, reviewed by experts, validated for compatibility, executed by predefined local skills and command-line tools, and exported with final outputs, intermediate step outputs, and audit records.",
        width=6.5,
    )

    doc.add_heading("Results", level=2)
    doc.add_heading("System-level workflow control enabled inspectable analysis planning", level=3)
    add_para(
        doc,
        "We first evaluated the agent as an engineering system rather than as a set of independent case studies. Each user request was converted into a structured workflow containing ordered steps, skill names, input references, parameters, expected outputs, and reasons for each operation. This representation made the analysis plan inspectable before execution. It also allowed a validator to test whether later steps referred to valid earlier outputs and whether the proposed skill was compatible with the uploaded file type.",
    )
    add_para(
        doc,
        "This design directly addressed common failure modes observed during development. Invalid workflows were most often caused by ambiguous natural-language prompts, incorrect source-file references, incompatible file formats, or hallucinated step-output names. The validator converted these failure modes into visible errors rather than allowing the local execution stage to fail silently. The resulting system therefore positioned the language model as a planner, not as an uncontrolled executor.",
    )
    add_compact_table(
        doc,
        ["Layer", "Risk addressed", "Control used", "Expert intervention"],
        [
            ["Planner", "Ambiguous user intent", "Structured workflow JSON", "Review proposed steps before running"],
            ["Validator", "Missing outputs or wrong file type", "Reference and compatibility checks", "Reject or revise invalid workflow"],
            ["Executor", "Uncontrolled code generation", "Predefined skills and local tools", "Approve local execution"],
            ["Reporter", "Opaque outputs", "Final files, step files, records and checksums", "Interpret results and boundaries"],
        ],
        widths=[1.2, 1.7, 1.8, 1.8],
    )

    doc.add_heading("Engineering controls targeted reproducibility failure modes", level=3)
    add_para(
        doc,
        "The agent was evaluated using engineering criteria that are important for ancient-data analysis: workflow validity, input-format compatibility, source metadata agreement, output traceability, and transparent handling of missing software environments. These criteria were selected because they describe whether a computational analysis can be reviewed and repeated, rather than whether a particular biological conclusion is true.",
    )
    add_para(
        doc,
        "For sequencing inputs, the most direct reproducibility checks were read-count agreement and file-size agreement between local files and public ENA/SRA records. Derived QC values such as GC content and mean read length were treated as descriptive outputs. This distinction prevented the agent from overclaiming: source-level agreement verifies that the correct files were processed, whereas biological interpretation remains the responsibility of domain experts.",
    )
    add_compact_table(
        doc,
        ["Evaluation target", "Metric or rule", "Interpretation"],
        [
            ["Workflow validity", "All step inputs must reference uploaded files or earlier outputs", "Prevents hallucinated pipeline references"],
            ["File compatibility", "Input format must match skill requirements", "Prevents applying FASTQ operations to CSV files"],
            ["Source agreement", "Δreads = 0 and Rbytes = 1 when public metadata are available", "Confirms local files match repository records"],
            ["Traceability", "Final outputs plus step outputs and run records", "Allows later inspection by users or supervisors"],
            ["Dependency behavior", "Missing tools are reported instead of auto-installed", "Keeps execution transparent on user machines"],
        ],
        widths=[1.6, 2.5, 2.4],
    )

    doc.add_heading("Case study 1: peptide CSV curation verified tabular skill behavior", level=3)
    add_para(
        doc,
        "As a lightweight tabular case, we tested a headerless peptide validation CSV containing 5,810 rows. The agent normalized the input into canonical label and sequence columns, validated peptide sequences, removed duplicate entries, calculated descriptive summaries, generated plots, and exported a cleaned CSV. The workflow produced 5,696 non-redundant peptide records and removed 114 duplicate sequence entries.",
    )
    add_para(
        doc,
        "This case was used to verify that the agent could handle ambiguous biological tables without manual spreadsheet editing. The output preserved the expected label structure and peptide-length distribution while producing a reusable cleaned table. Because peptide analysis is only one possible application, this result is presented as a skill-behavior check rather than as the central biological contribution of the system.",
    )
    add_figure(
        doc,
        TASK_DIR / "task1_peptide_csv_curation.png",
        "Figure 2. Peptide CSV curation. A, Natural-language peptide-table processing workflow. B, Row-count change after duplicate removal. C, Label balance before and after cleaning. D, Peptide-length distribution after curation.",
        width=6.5,
    )

    doc.add_heading("Case study 2: FASTQ QC verified source-level rigor for sequencing inputs", level=3)
    add_para(
        doc,
        "We then evaluated six public FASTQ records relevant to ancient or paleomicrobiome source-data reuse. The agent performed file inspection, FastQC, MultiQC, seqkit-based statistics, and metadata comparison against ENA/SRA records. Across all six records, agent-derived read counts matched the corresponding public records with Δreads = 0. Local compressed file sizes also matched public fastq_bytes values, giving Rbytes = 1.000 for all tested inputs. The tested files ranged from 56,122 to 4,380,359 reads.",
    )
    add_para(
        doc,
        "These results show that the agent can verify whether local sequencing files correspond to the intended public records before downstream analysis. For ancient DNA and paleomicrobiome work, this is a necessary front-end control because incorrect source files, incomplete downloads, or paired-end mismatches can invalidate later alignment, taxonomic, or authentication analyses even when downstream tools run successfully.",
    )
    add_figure(
        doc,
        TASK_DIR / "task2_fastq_source_qc.png",
        "Figure 3. FASTQ source-data quality control. A, Agent-derived read counts compared with public ENA/SRA read counts. B, Local compressed file sizes compared with public fastq_bytes values. C, GC content reported as a derived QC metric. D, Source-record agreement summarized across the tested files.",
        width=6.5,
    )

    doc.add_heading("Case study 3: retraction-associated audit tested boundary-aware analysis", level=3)
    add_para(
        doc,
        "Finally, we used SRR29088443 from SRP508771 as a negative-control audit case because the associated publication record is linked to a retraction. The goal was not to evaluate the retracted biological claim, but to test whether the agent could audit public source data from a problematic literature context while clearly limiting the conclusion to file-level reproducibility.",
    )
    add_para(
        doc,
        "The agent identified the paired FASTQ files and reproduced the public metadata. Each mate contained 53,571 reads. The total base count was 26,517,645 bp, and compressed file sizes matched the public record for both R1 and R2. Derived QC outputs included mean read lengths of 244 bp and 251 bp, GC contents of 54.0% and 53.6%, and zero malformed records. This case demonstrates that the system can document source-data consistency while explicitly avoiding unsupported claims about a contested biological interpretation.",
    )
    add_figure(
        doc,
        RETRACT_DIR / "task3_retracted_fastq_audit.png",
        "Figure 4. Retraction-associated source-data audit. A, Public record and accession context. B, Agent-derived metadata deviation from public records. C, Derived FASTQ QC metrics. D, Audit checklist separating source-data verification from biological conclusion validation.",
        width=6.5,
    )

    doc.add_heading("Discussion", level=2)
    add_para(
        doc,
        "This study presents a prototype human-in-the-loop research agent for rigorous early-stage paleomicrobiome and ancient-data analysis. The main contribution is not a new ancient-DNA algorithm, but a controlled natural-language interface that helps researchers move from expert intent to validated, locally executed, and auditable workflows. This framing is important because many errors in ancient-data reuse occur before advanced biological interpretation begins.",
    )
    add_para(
        doc,
        "The results indicate that separating planning, validation, execution, and interpretation can make natural-language agents more suitable for scientific work. The language model provides flexibility when translating user intent, while predefined skills and validators restrict what can be executed. Expert review points preserve human responsibility for task definition and interpretation. This structure is particularly relevant for paleomicrobiome studies, where the same file-level operation may have different implications depending on sample preservation, contamination risk, laboratory context, and downstream authentication criteria.",
    )
    add_para(
        doc,
        "The current prototype should be interpreted as an auditable front end rather than a complete ancient-DNA analysis platform. It supports file-level checks, table curation, quality-control report generation, and source-record comparison. Future versions should extend the skill library to alignment, duplicate marking, damage-pattern estimation, contamination screening, taxonomic profiling, and authentication reporting. These extensions should retain the same human-in-the-loop principle: the system may execute standardized operations, but experts must define the question, approve the workflow, and interpret biological meaning.",
    )
    add_para(
        doc,
        "For practical usability, the tool also requires a public code repository and a usable web interface. The present implementation already uses a local browser-based interface for file upload, natural-language workflow generation, local-execution approval, progress display, and organized output export. Before submission or public distribution, the code should be released on GitHub with example datasets, installation instructions, environment checks, and a stable release tag. A hosted demonstration website or packaged desktop release would further improve accessibility for non-programming users, while preserving local execution for sensitive or large files.",
    )

    doc.add_heading("Methods", level=2)
    doc.add_heading("Agent workflow representation", level=3)
    add_para(
        doc,
        "A user request was represented as an ordered workflow W containing n executable steps:",
    )
    add_formula(doc, "W = {s₁, s₂, …, sₙ}")
    add_para(
        doc,
        "Each step sᵢ was defined by a skill identifier kᵢ, input references Iᵢ, output declarations Oᵢ, parameters θᵢ, and a natural-language reason rᵢ:",
    )
    add_formula(doc, "sᵢ = (kᵢ, Iᵢ, Oᵢ, θᵢ, rᵢ)")
    add_para(
        doc,
        "This representation made the planned analysis explicit before execution and allowed both automated validation and human review.",
    )

    doc.add_heading("Workflow validation", level=3)
    add_para(
        doc,
        "The validator evaluated a proposed workflow W against the available uploaded files F and the available skill registry K:",
    )
    add_formula(doc, "V(W, F, K) → {0, 1}")
    add_para(
        doc,
        "A workflow was accepted only when every step used an available skill, every input reference pointed to an uploaded file or an earlier step output, and every input format was compatible with the requested skill:",
    )
    add_formula(doc, "Iᵢ ⊆ F ∪ O<ᵢ")
    add_formula(doc, "format(Iᵢ) ∈ compatible(kᵢ)")
    add_para(
        doc,
        "These checks were added because unconstrained planning produced errors such as missing step outputs, inconsistent reference syntax, and inappropriate decompression or FASTQ-QC operations for the selected file type.",
    )

    doc.add_heading("Execution environment and implemented skills", level=3)
    add_para(
        doc,
        "Validated steps were dispatched to local skills. Implemented skills included peptide CSV normalization, peptide sequence validation, duplicate removal, CSV export, file-type detection, paired-end FASTQ checking, FastQC execution, MultiQC report generation, seqkit statistics, and source-data audit utilities. External tools were called from configured local environments. If a required program was missing, the system returned an explicit missing-dependency message rather than installing software automatically.",
    )

    doc.add_heading("Human-in-the-loop checkpoints", level=3)
    add_para(
        doc,
        "Human intervention was built into three checkpoints. First, the expert reviewed whether the generated workflow matched the scientific request. Second, the user approved local execution after inspecting the planned steps. Third, the user interpreted the final reports with access to intermediate step outputs and audit records. These checkpoints were designed to keep the agent useful for non-programming users without transferring scientific judgment to the model.",
    )

    doc.add_heading("Source-level agreement metrics", level=3)
    add_para(
        doc,
        "For FASTQ source-data verification, read-count and compressed-file-size agreement were calculated relative to public ENA/SRA metadata:",
    )
    add_formula(doc, "Δreads = Nagent − Npublic")
    add_formula(doc, "Rbytes = Blocal / Bpublic")
    add_para(
        doc,
        "A local FASTQ record was considered source-metadata matched when Δreads = 0 and Rbytes = 1 for the relevant public record. These metrics assessed whether the local file matched the repository source record, not whether downstream biological claims were correct.",
    )

    doc.add_heading("FASTQ quality-control summaries", level=3)
    add_para(
        doc,
        "Derived FASTQ summaries included GC percentage, mean read length, and the fraction of ambiguous bases. For a FASTQ file with N reads and read lengths Lⱼ, the reported summaries were:",
    )
    add_formula(doc, "GC(%) = (G + C) / (A + T + G + C + N) × 100")
    add_formula(doc, "L̄ = (1/N) Σⱼ Lⱼ")
    add_formula(doc, "fN = Nbases,N / Nbases,total")
    add_para(
        doc,
        "FastQC and MultiQC reports were generated as standard QC artifacts. Seqkit-derived values were used for compact source-level summaries and figure preparation.",
    )

    doc.add_heading("Dataset-level evaluations", level=3)
    add_para(
        doc,
        "The peptide CSV case evaluated tabular curation behavior using row counts, duplicate counts, unique sequence counts, label distribution, and peptide-length distribution. The public FASTQ case evaluated six records using read-count agreement, compressed-file-size agreement, GC content, and mean read length. The retraction-associated case evaluated SRR29088443 as a boundary test: source-data reproducibility was audited, but the biological interpretation of the associated retracted article was not assessed.",
    )

    doc.add_heading("Output organization and web interface", level=3)
    add_para(
        doc,
        "Each local run generated a timestamped result directory containing final_outputs for user-requested files, step_outputs for intermediate files, and ResearchAgent Records for workflow JSON, manifests, checksums, timestamps, and run reports. The local browser interface allowed users to upload files, enter a natural-language task, inspect the generated workflow, approve local execution, monitor progress, and open the final report. This web-interface layer is central to the human-in-the-loop design because it makes intermediate decisions visible to expert users.",
    )

    doc.add_heading("Data Availability", level=2)
    add_para(
        doc,
        "All validation datasets used in this study are public or organized as local example files in the project workspace. Public sequencing records were retrieved from ENA/SRA using the accession numbers listed below. Agent-generated cleaned outputs, quality-control reports, workflow JSON files, intermediate step outputs, and run manifests are stored in the local result directories and should be deposited with the code repository before release.",
    )
    add_compact_table(
        doc,
        ["Dataset group", "Data type", "Source or accession", "Role in this study", "Availability"],
        [
            [
                "Peptide validation table",
                "CSV peptide table",
                "Validation.csv; local validation file derived from the peptide dataset used in this study",
                "Tabular curation, sequence validation, duplicate removal and export test",
                "Project example data; source-paper details to be finalized before submission",
            ],
            [
                "FASTQ source-QC set",
                "Public FASTQ records",
                "ERR15682270, ERR10114877, ERR3250149, ERR10114867, ERR10114861 and ERR15682267",
                "Read-count agreement, file-size agreement and FASTQ QC evaluation",
                "ENA/SRA public records",
            ],
            [
                "Retraction-associated audit case",
                "Public paired-end FASTQ record",
                "SRR29088443 from SRP508771",
                "Boundary-aware source-data audit for a problematic literature record",
                "SRA/ENA public record",
            ],
        ],
        widths=[1.15, 1.05, 1.65, 1.75, 1.15],
        font_size=8,
    )
    doc.add_heading("Code Availability", level=2)
    add_para(
        doc,
        "The Research Agent code is currently maintained in the local project repository. Before public distribution, the repository should be released on GitHub with a stable release tag, installation guide, example datasets, environment-check instructions, and a minimal web-interface demonstration. If a hosted website is not feasible before submission, the manuscript should state that the tool is distributed as a local web application that runs on the user's machine.",
    )
    doc.add_heading("Acknowledgments", level=2)
    add_para(doc, "Acknowledgments to supervisors, collaborators, and funding sources should be completed by the author.")
    doc.add_heading("Author contributions", level=2)
    add_para(doc, "Author contribution statements should be completed after the final author list is confirmed.")
    doc.add_heading("Competing interest", level=2)
    add_para(doc, "The authors declare no competing interests.")
    add_refs(doc)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
