from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ancient_research_agent_manuscript_draft.docx"
FIG_DIR = ROOT / "Final data" / "manuscript_figures"
TASK_DIR = ROOT / "Final data" / "task_grouped_figures"
RETRACT_DIR = ROOT / "Final data" / "retracted_case_figure"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_reference_table(doc: Document) -> None:
    doc.add_heading("Manuscript working checklist", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Component", "Current status", "Author notes"]):
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, text, bold=True)
    rows = [
        ("Graphical abstract", "Inserted", "Three application scenarios: reproduce, check, audit."),
        ("Figure 1", "Inserted", "System architecture and constrained local execution."),
        ("Figure 2", "Inserted as two panels", "Task 1 peptide CSV and Task 2 FASTQ source-data QC."),
        ("Figure 4", "Inserted", "Task 3: retraction-associated FASTQ audit case."),
        ("References", "Draft placeholders", "Replace or expand with final citation manager output."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)


def set_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2F6F8F"),
        ("Heading 3", 12, "1F2933"),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(12 if name != "Heading 3" else 8)
        st.paragraph_format.space_after = Pt(6 if name != "Heading 3" else 4)

    if "Caption" in styles:
        cap = styles["Caption"]
        cap.font.name = "Calibri"
        cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        cap.font.size = Pt(9)
        cap.font.italic = False
        cap.font.color.rgb = RGBColor.from_string("4B5563")


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph(style="Heading 1")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("A natural-language research agent for reproducible ancient human and paleomicrobiome data processing")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("1F2933")

    for text in [
        "Author names to be completed",
        "Affiliations to be completed",
        "# These authors contributed equally, if applicable.",
        "*Correspondence should be addressed to: [corresponding author email]",
        "Keywords: ancient DNA; paleomicrobiome; research agent; reproducibility; FASTQ quality control; source-data audit",
    ]:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(3)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Figure file missing: {path.name}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(caption)


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_styles(doc)

    add_title_block(doc)

    doc.add_heading("Abstract", level=2)
    add_para(
        doc,
        "Ancient human and paleomicrobiome studies increasingly rely on public sequencing repositories and heterogeneous supplementary data, yet practical reuse remains difficult because users must identify source files, validate formats, run specialized command-line tools, and preserve intermediate records. Here, we present a prototype natural-language research agent designed to support reproducible entry-level processing of ancient-data-related resources. The agent converts user requests into structured workflows, validates step compatibility, dispatches predefined local skills and command-line tools, and exports final outputs together with intermediate files and audit records. We evaluated the prototype using three representative scenarios: peptide validation table curation, FASTQ source-data quality control against public ENA/SRA metadata, and a retraction-associated sequencing record used as a negative-control audit case. Across these tasks, the agent organized heterogeneous inputs, reproduced source-level metadata, and generated traceable quality-control outputs. This framework does not replace expert interpretation, but provides a controlled natural-language interface for transparent and auditable biological data reuse.",
    )

    doc.add_heading("Key points", level=2)
    add_bullets(
        doc,
        [
            "A natural-language agent was developed to convert user requests into constrained, locally executable research workflows.",
            "The prototype focuses on entry-level but failure-prone steps in ancient human and paleomicrobiome data reuse, including source-data checks, FASTQ QC, paired-end consistency, and audit-trail generation.",
            "Validation tasks covered peptide CSV curation, public FASTQ metadata reproduction, and a retraction-associated FASTQ source-data audit.",
            "The language model is used for workflow planning, whereas predefined skills and local tools perform deterministic execution.",
        ],
    )

    doc.add_heading("Graphical abstract", level=2)
    add_figure(
        doc,
        FIG_DIR / "graphical_abstract_three_scenarios.png",
        "Graphical abstract. The agent supports three practical scenarios in ancient human and paleomicrobiome data reuse: routine reproducibility, source-data checking, and audit of problematic or retraction-associated records.",
        width=6.5,
    )

    doc.add_heading("Introduction", level=2)
    add_para(
        doc,
        "Ancient human, ancient DNA, and paleomicrobiome studies depend heavily on public data reuse. Although many studies provide accessions, supplementary tables, and source files, reproducing even the early processing steps can be technically demanding. Users must locate the correct files, check whether local downloads match repository records, run command-line tools, and preserve outputs in a form that can be inspected later. These tasks are especially important for ancient-data applications because downstream conclusions often depend on fragile, low-input, degraded, or contamination-sensitive source material [1–3].",
    )
    add_para(
        doc,
        "For ancient DNA and paleomicrobiome datasets, entry-level checks are not merely administrative. FASTQ file integrity, paired-end consistency, read-count agreement, read-length distributions, GC content, and missing or malformed records all affect the reliability of subsequent mapping, damage profiling, taxonomic profiling, and contamination assessment. If the wrong source file is used, if paired files are mismatched, or if repository metadata are not checked, downstream analyses may be compromised before biological interpretation begins.",
    )
    add_para(
        doc,
        "Large language models provide a promising interface for lowering the technical barrier to data processing, but unrestricted code generation is not sufficiently stable for reproducible research. Workflow hallucination, incompatible file references, and missing tool environments can all lead to silent or confusing failures. We therefore designed a research agent in which the language model proposes a structured workflow, while predefined skills, validators, and local tools control execution. This design separates flexible natural-language understanding from deterministic computational operations.",
    )
    add_para(
        doc,
        "In this study, we describe a prototype natural-language research agent for reproducible ancient human and paleomicrobiome data processing. The current implementation emphasizes critical entry points in data reuse rather than full downstream biological interpretation. We demonstrate the system using table curation, FASTQ source-data quality control, and retraction-associated source-data auditing.",
    )

    add_figure(
        doc,
        FIG_DIR / "figure1_agent_architecture_workflow.png",
        "Figure 1. System architecture of the natural-language research agent. User requests are translated into structured workflows, checked by a validator, executed by predefined local skills and command-line tools, and exported as final outputs, step outputs, and audit records.",
        width=6.5,
    )

    doc.add_heading("Results", level=2)
    doc.add_heading("Task 1: The agent curated peptide CSV data into a reproducible validation table", level=3)
    add_para(
        doc,
        "We first evaluated whether the agent could process a small tabular biological dataset. The input was a peptide validation CSV file with no explicit header. From a natural-language request, the agent generated a workflow to normalize the table, identify label and sequence columns, validate peptide sequences, remove duplicate entries, compute summary statistics, and export a cleaned CSV file. This task tested whether the agent could convert an ambiguous table into a standardized downstream-ready dataset while preserving interpretable audit records.",
    )
    add_para(
        doc,
        "The resulting workflow successfully normalized the input into label and sequence fields and produced a non-redundant validation table. Duplicate sequence removal reduced the table size while maintaining the expected label structure and peptide length distribution. This demonstrates that the agent can perform constrained curation of biological tabular data rather than relying on ad hoc manual spreadsheet operations.",
    )
    add_figure(
        doc,
        TASK_DIR / "task1_peptide_csv_curation.png",
        "Figure 2. Peptide CSV curation. The agent detects and normalizes a headerless peptide table, validates peptide sequences, removes duplicate records, and exports a cleaned validation table while preserving key distributional properties.",
        width=6.5,
    )

    doc.add_heading("Task 2: The agent reproduced public FASTQ source metadata", level=3)
    add_para(
        doc,
        "We next tested public FASTQ source-data quality control. In this task, the agent processed sequencing files using FastQC, MultiQC, and seqkit, then compared agent-derived metrics with ENA/SRA public records. The comparison focused on source-level properties such as read counts and file sizes, while GC content and mean read length were treated as derived quality-control metrics rather than claims to be matched to a biological result.",
    )
    add_para(
        doc,
        "Across the tested FASTQ records, agent-derived read counts matched the corresponding public records, and local file sizes matched public repository metadata. These results indicate that the agent can support a key early step in ancient-data reuse: confirming that local inputs correspond to the intended public sequencing records before downstream analysis.",
    )
    add_figure(
        doc,
        TASK_DIR / "task2_fastq_source_qc.png",
        "Figure 3. FASTQ source-data quality control. Agent-derived sequencing metrics are compared with public ENA/SRA records, while GC content and read-length summaries are reported as quality-control outputs.",
        width=6.5,
    )

    doc.add_heading("Task 3: The agent audited a retraction-associated public sequencing record", level=3)
    add_para(
        doc,
        "Finally, we evaluated the agent on a retraction-associated public sequencing record as a negative-control audit case. The selected paired FASTQ run was associated with a retracted publication but was used here only to test source-data processing and metadata reproduction. The agent was not used to validate the original biological conclusion or infer the reason for retraction.",
    )
    add_para(
        doc,
        "For this case, the agent identified the paired FASTQ files, reproduced public read-count, base-count, and file-size metadata, and generated FastQC, MultiQC, and seqkit outputs. This result illustrates a distinct use case: transparent audit of source files linked to problematic literature records. The audit confirms whether the public source data can be processed consistently, while preserving a clear boundary between source-data verification and biological interpretation.",
    )
    add_figure(
        doc,
        RETRACT_DIR / "task3_retracted_fastq_audit.png",
        "Figure 4. Task 3: retraction-associated FASTQ audit. The agent reproduces public source-level metadata for a retraction-associated sequencing record and generates QC outputs, while explicitly avoiding claims about the retracted biological conclusion.",
        width=6.5,
    )

    doc.add_heading("Discussion", level=2)
    add_para(
        doc,
        "This prototype shows that a natural-language interface can be combined with constrained local execution to support reproducible biological data reuse. The main contribution is not a new ancient-DNA algorithm, but an execution framework that helps users move from natural-language intent to organized, inspectable, and locally generated outputs. This is particularly relevant for ancient human and paleomicrobiome studies, where early source-data mistakes can propagate into downstream analyses.",
    )
    add_para(
        doc,
        "The agent differs from a general-purpose chatbot because it does not directly execute arbitrary generated code. Instead, it plans workflows that must be compatible with predefined skills and validated file relationships. This structure reduces the risk of hallucinated steps, incompatible file references, and opaque execution. The separation between final outputs, step outputs, and research-agent records also supports later inspection and reproducibility.",
    )
    add_para(
        doc,
        "Several limitations remain. The current prototype focuses on entry-level source-data processing rather than full ancient-DNA pipelines. Downstream analyses such as read alignment, damage-pattern estimation, contamination assessment, metagenomic classification, and authentication remain future extensions. The quality of generated workflows also depends on the planner prompt, validator coverage, and skill definitions. In addition, external tools still need to be available in the local environment, although the agent can report missing dependencies rather than silently installing them.",
    )
    add_para(
        doc,
        "Future work will extend the skill library toward ancient-DNA-specific analysis modules, including BWA or Bowtie2 alignment, samtools statistics, mapDamage profiling, contamination screening, and taxonomic profiling. A larger benchmark across published paleogenomic and paleomicrobiome datasets will also be needed to evaluate robustness across data types, repository conventions, and laboratory contexts.",
    )

    doc.add_heading("Methods", level=2)
    doc.add_heading("Agent workflow planning", level=3)
    add_para(
        doc,
        "The agent accepts uploaded data files and a natural-language user request. The planner converts the request into a structured workflow containing ordered steps, each specifying a skill name, inputs, parameters, expected outputs, and a textual reason. This representation allows subsequent validation before local execution.",
    )
    doc.add_heading("Workflow validation and local skill execution", level=3)
    add_para(
        doc,
        "Before execution, the workflow validator checks file-format compatibility, step references, and expected output relationships. Validated steps are dispatched to local skills. Current skills include peptide CSV normalization and validation, FASTQ file-type detection, paired-end matching, FastQC execution, MultiQC summary generation, seqkit statistics, and source-data audit helpers.",
    )
    doc.add_heading("Output organization and audit records", level=3)
    add_para(
        doc,
        "Each run creates an independent result directory. Final user-facing outputs are exported to a final_outputs folder, while intermediate files are stored in step_outputs. Workflow JSON, manifests, checksums, timestamps, and an HTML report are preserved in ResearchAgent Records. This organization separates user-facing deliverables from reproducibility evidence.",
    )
    doc.add_heading("Validation datasets", level=3)
    add_para(
        doc,
        "Validation tasks were selected to represent common entry points in ancient human and paleomicrobiome data reuse rather than full biological conclusion reproduction. The peptide CSV task assessed heterogeneous table curation. The FASTQ source-data QC task assessed agreement between local files and public repository metadata. The retraction-associated task assessed whether the agent could transparently process and audit a problematic literature-linked sequencing record without making biological claims.",
    )

    doc.add_heading("Data Availability", level=2)
    add_para(
        doc,
        "The validation datasets used in this draft are public or locally organized within the project workspace. Public sequencing records are linked to ENA/SRA accessions described in the manuscript. Final accession lists and dataset download commands should be added before submission.",
    )

    doc.add_heading("Code Availability", level=2)
    add_para(
        doc,
        "The Research Agent code is currently maintained in the local project repository. A public GitHub repository, installation instructions, example data, and release tag should be added before submission or public distribution.",
    )

    doc.add_heading("Acknowledgments", level=2)
    add_para(doc, "Acknowledgments to supervisors, collaborators, and funding sources should be completed by the author.")

    doc.add_heading("Author contributions", level=2)
    add_para(doc, "Author contribution statements should be completed after the final author list is confirmed.")

    doc.add_heading("Competing interest", level=2)
    add_para(doc, "The authors declare no competing interests.")

    doc.add_heading("References", level=1)
    refs = [
        "1. Pääbo, S. Molecular cloning of Ancient Egyptian mummy DNA. Nature 314, 644–645 (1985).",
        "2. Dabney, J., Meyer, M. & Pääbo, S. Ancient DNA damage. Cold Spring Harb. Perspect. Biol. 5, a012567 (2013).",
        "3. Warinner, C., et al. A robust framework for microbial archaeology. Annu. Rev. Genomics Hum. Genet. 18, 321–356 (2017).",
        "4. Andrews, S. FastQC: a quality control tool for high throughput sequence data (2010).",
        "5. Ewels, P., Magnusson, M., Lundin, S. & Käller, M. MultiQC: summarize analysis results for multiple tools and samples in a single report. Bioinformatics 32, 3047–3048 (2016).",
        "6. Shen, W., Le, S., Li, Y. & Hu, F. SeqKit: a cross-platform and ultrafast toolkit for FASTA/Q file manipulation. PLoS ONE 11, e0163962 (2016).",
        "7. Leinonen, R., Sugawara, H. & Shumway, M. The Sequence Read Archive. Nucleic Acids Res. 39, D19–D21 (2011).",
        "8. Harrison, P. W., et al. The European Nucleotide Archive in 2024. Nucleic Acids Res. 52, D92–D97 (2024).",
        "9. Schubert, M., Lindgreen, S. & Orlando, L. AdapterRemoval v2: rapid adapter trimming, identification, and read merging. BMC Res. Notes 9, 88 (2016).",
        "10. Jónsson, H., Ginolhac, A., Schubert, M., Johnson, P. L. F. & Orlando, L. mapDamage2.0: fast approximate Bayesian estimates of ancient DNA damage parameters. Bioinformatics 29, 1682–1684 (2013).",
        "11. Gao, W., et al. Predictable regulation of gut microbiome in immunotherapeutic efficacy of gastric cancer. Genes Immun. 26, 1–8 (2025). [Retracted]",
        "12. Gao, W., et al. Retraction Note: Predictable regulation of gut microbiome in immunotherapeutic efficacy of gastric cancer. Genes Immun. 27, 384 (2026).",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    doc.add_page_break()
    add_reference_table(doc)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
